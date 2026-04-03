from __future__ import annotations

from datetime import timedelta
from io import BytesIO

from django.contrib import admin, messages
from django.contrib.auth.admin import UserAdmin as DjangoUserAdmin
from django.core.exceptions import PermissionDenied, ValidationError
from django.db.models import Count, Q, Sum
from django.forms import ModelForm
from django.http import HttpRequest, HttpResponse
from django.urls import path, reverse
from django.utils import timezone
from django.utils.html import format_html

from crm.admin_views import cycle_finance_summary_view, teacher_performance_view
from crm.models import (
    AttendanceRecord,
    BillingCycle,
    GroupEnrollment,
    GroupSchedule,
    Lesson,
    Payment,
    PlacementTest,
    StudyGroup,
    Subject,
    TestAttempt,
    User,
    UserExitRecord,
)

try:
    from docx import Document
except ImportError:  # pragma: no cover
    Document = None

try:
    from openpyxl import Workbook
except ImportError:  # pragma: no cover
    Workbook = None

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
except ImportError:  # pragma: no cover
    A4 = None
    canvas = None

try:
    from unfold.admin import ModelAdmin
except ImportError:  # pragma: no cover
    from django.contrib.admin import ModelAdmin


def is_super_admin(user: User) -> bool:
    return bool(user.is_authenticated and (user.is_superuser or user.role == User.Role.SUPER_ADMIN))


def is_center_admin(user: User) -> bool:
    return bool(user.is_authenticated and user.role == User.Role.CENTER_ADMIN)


def is_teacher(user: User) -> bool:
    return bool(user.is_authenticated and user.role == User.Role.TEACHER)


def is_student(user: User) -> bool:
    return bool(user.is_authenticated and user.role == User.Role.STUDENT)


class UserExitRecordForm(ModelForm):
    class Meta:
        model = UserExitRecord
        fields = "__all__"

    def clean_reason(self):
        reason = (self.cleaned_data.get("reason") or "").strip()
        if not reason:
            raise ValidationError("Причина выбытия обязательна.")
        return reason


class TestAttemptInline(admin.TabularInline):
    model = TestAttempt
    extra = 0
    fields = ("test", "total_questions", "correct_answers", "score_percent", "level", "is_final", "taken_at")
    readonly_fields = ("score_percent", "level", "taken_at")

    def has_module_permission(self, request):
        return is_super_admin(request.user) or is_center_admin(request.user)

    def has_add_permission(self, request, obj=None):
        return is_super_admin(request.user) or is_center_admin(request.user)

    def has_change_permission(self, request, obj=None):
        return is_super_admin(request.user) or is_center_admin(request.user)

    def has_delete_permission(self, request, obj=None):
        return is_super_admin(request.user)


@admin.register(User)
class UserAdmin(DjangoUserAdmin):
    conditional_fields = {}

    list_display = (
        "username",
        "full_name",
        "role",
        "placement_level",
        "is_active_in_center",
        "is_staff",
        "is_active",
    )
    list_filter = ("role", "placement_level", "is_active_in_center", "is_staff", "is_active")
    search_fields = (
        "username",
        "first_name",
        "last_name",
        "middle_name",
        "phone_student",
        "phone_parent",
    )
    ordering = ("username",)
    inlines = [TestAttemptInline]

    readonly_fields = (
        "photo_preview",
        "latest_test_summary",
        "student_groups_summary",
        "student_progress_last_3_months",
        "teacher_workload_summary",
        "export_user_pdf_button",
    )

    fieldsets = DjangoUserAdmin.fieldsets + (
        (
            "Профиль центра",
            {
                "fields": (
                    "role",
                    "middle_name",
                    "phone_student",
                    "phone_parent",
                    "address",
                    "photo",
                    "photo_preview",
                    "school_name",
                    "school_shift",
                    "teacher_specialization",
                    "teacher_experience_years",
                    "joined_center_at",
                    "is_active_in_center",
                )
            },
        ),
        (
            "Тестирование уровня",
            {
                "fields": (
                    "placement_score_percent",
                    "placement_level",
                    "latest_test_summary",
                )
            },
        ),
    )

    add_fieldsets = DjangoUserAdmin.add_fieldsets + (
        (
            "Профиль центра",
            {
                "fields": (
                    "role",
                    "middle_name",
                    "phone_student",
                    "phone_parent",
                    "address",
                    "photo",
                    "school_name",
                    "school_shift",
                    "teacher_specialization",
                    "teacher_experience_years",
                )
            },
        ),
    )

    def _filter_fieldsets(self, fieldsets, forbidden_fields: set[str]):
        filtered = []
        for section, options in fieldsets:
            fields = options.get("fields", ())
            updated_fields = []
            for field in fields:
                if isinstance(field, (tuple, list)):
                    nested = tuple(item for item in field if item not in forbidden_fields)
                    if nested:
                        updated_fields.append(nested)
                elif field not in forbidden_fields:
                    updated_fields.append(field)

            if updated_fields:
                filtered.append((section, {**options, "fields": tuple(updated_fields)}))
        return tuple(filtered)

    def get_fieldsets(self, request, obj=None):
        fieldsets = list(super().get_fieldsets(request, obj))

        if obj and obj.role == User.Role.STUDENT:
            fieldsets.append(
                (
                    "Панель студента",
                    {
                        "fields": (
                            "student_groups_summary",
                            "student_progress_last_3_months",
                            "export_user_pdf_button",
                        )
                    },
                )
            )
        elif obj and obj.role == User.Role.TEACHER:
            fieldsets.append(
                (
                    "Панель учителя",
                    {
                        "fields": (
                            "teacher_workload_summary",
                            "export_user_pdf_button",
                        )
                    },
                )
            )
        elif obj:
            fieldsets.append(("Экспорт", {"fields": ("export_user_pdf_button",)}))

        forbidden_fields = {"groups", "user_permissions"}

        # Students and teachers should never manually touch global permissions.
        if is_center_admin(request.user) or is_teacher(request.user) or is_student(request.user):
            forbidden_fields.update({"is_superuser"})

        return self._filter_fieldsets(tuple(fieldsets), forbidden_fields)

    def get_readonly_fields(self, request, obj=None):
        readonly = list(super().get_readonly_fields(request, obj))
        readonly.extend(self.readonly_fields)

        if is_center_admin(request.user):
            readonly.extend(["is_staff", "is_superuser", "role", "placement_score_percent", "placement_level"])

        if (is_teacher(request.user) or is_student(request.user)) and obj is not None:
            # Own profile is view-only for teacher/student.
            model_fields = [field.name for field in self.model._meta.fields]
            readonly.extend(model_fields)

        return tuple(dict.fromkeys(readonly))

    def get_form(self, request, obj=None, **kwargs):
        form = super().get_form(request, obj, **kwargs)
        if "role" in form.base_fields and is_center_admin(request.user):
            form.base_fields["role"].choices = [
                (User.Role.TEACHER, User.Role.TEACHER.label),
                (User.Role.STUDENT, User.Role.STUDENT.label),
            ]
        return form

    def get_queryset(self, request):
        queryset = super().get_queryset(request)
        if is_super_admin(request.user):
            return queryset
        if is_center_admin(request.user):
            return queryset.filter(role__in=[User.Role.TEACHER, User.Role.STUDENT], exit_record__isnull=True)
        if is_teacher(request.user) or is_student(request.user):
            return queryset.filter(pk=request.user.pk)
        return queryset.none()

    def has_module_permission(self, request):
        return is_super_admin(request.user) or is_center_admin(request.user) or is_teacher(request.user) or is_student(
            request.user
        )

    def has_view_permission(self, request, obj=None):
        if is_super_admin(request.user) or is_center_admin(request.user):
            return True
        if is_teacher(request.user) or is_student(request.user):
            return obj is None or obj.pk == request.user.pk
        return False

    def has_add_permission(self, request):
        return is_super_admin(request.user) or is_center_admin(request.user)

    def has_change_permission(self, request, obj=None):
        if is_super_admin(request.user) or is_center_admin(request.user):
            return True
        if is_teacher(request.user) or is_student(request.user):
            return obj is None or obj.pk == request.user.pk
        return False

    def has_delete_permission(self, request, obj=None):
        return is_super_admin(request.user)

    def save_model(self, request, obj, form, change):
        is_new = obj.pk is None
        if is_center_admin(request.user) and obj.role not in {User.Role.TEACHER, User.Role.STUDENT}:
            raise PermissionDenied("Администрация центра может создавать только Учителей и Студентов.")

        if obj.role == User.Role.STUDENT:
            obj.is_staff = True
            obj.is_superuser = False
        elif obj.role == User.Role.TEACHER:
            obj.is_staff = True
            obj.is_superuser = False
        elif obj.role == User.Role.CENTER_ADMIN:
            obj.is_staff = True
            obj.is_superuser = False
        elif obj.role == User.Role.SUPER_ADMIN:
            obj.is_staff = True
            obj.is_superuser = True

        if obj.role in {User.Role.STUDENT, User.Role.TEACHER} and is_new:
            obj.is_active_in_center = False
            obj.is_active = False

        super().save_model(request, obj, form, change)

        if obj.role in {User.Role.STUDENT, User.Role.TEACHER, User.Role.CENTER_ADMIN}:
            obj.groups.clear()
            obj.user_permissions.clear()

        if obj.role in {User.Role.STUDENT, User.Role.TEACHER}:
            has_final_attempt = TestAttempt.objects.filter(user_id=obj.pk, is_final=True).exists()
            if not has_final_attempt and (obj.is_active_in_center or obj.is_active):
                obj.is_active_in_center = False
                obj.is_active = False
                obj.save(update_fields=["is_active_in_center", "is_active"])
            if not has_final_attempt:
                self.message_user(
                    request,
                    "Пользователь сохранен, но не активирован в центре до завершения тестирования.",
                    level=messages.WARNING,
                )

    @admin.display(description="Фото")
    def photo_preview(self, obj: User):
        if obj and obj.photo:
            return format_html(
                '<img src="{}" style="max-height:120px;border-radius:8px;border:1px solid #d0d5dd;"/>',
                obj.photo.url,
            )
        return "Фото не загружено"

    @admin.display(description="Последний результат теста")
    def latest_test_summary(self, obj: User):
        attempt = obj.latest_final_attempt
        if not attempt:
            return "Тест не пройден"
        return format_html(
            "{}% ({}) от {}",
            attempt.score_percent,
            attempt.get_level_display(),
            attempt.taken_at.strftime("%Y-%m-%d %H:%M"),
        )

    @admin.display(description="Группы и расписание")
    def student_groups_summary(self, obj: User):
        if obj.role != User.Role.STUDENT:
            return "Только для роли Студент"

        enrollments = (
            GroupEnrollment.objects.filter(student=obj, is_active=True)
            .select_related("group__subject", "group__teacher")
            .prefetch_related("group__schedules", "group__enrollments__student")
        )
        if not enrollments.exists():
            return "Студент пока не добавлен в группы"

        chunks = []
        for enrollment in enrollments:
            group = enrollment.group
            schedules = ", ".join(
                f"{schedule.get_day_of_week_display()} {schedule.start_time:%H:%M}-{schedule.end_time:%H:%M}"
                for schedule in group.schedules.all()
            ) or "Нет расписания"

            classmates = ", ".join(
                student.full_name
                for student in group.enrollments.filter(is_active=True).exclude(student=obj).select_related("student")[:10]
                for student in [student.student]
            ) or "Нет данных"

            teacher_name = group.teacher.full_name if group.teacher else "Не назначен"
            teacher_phone = group.teacher.phone_student if group.teacher and group.teacher.phone_student else "-"

            chunks.append(
                f"<div style='margin-bottom:10px;'><strong>{group.subject.name} / {group.name}</strong><br/>"
                f"Учитель: {teacher_name} ({teacher_phone})<br/>"
                f"Расписание: {schedules}<br/>"
                f"Одногруппники: {classmates}</div>"
            )

        return format_html("".join(chunks))

    @admin.display(description="Успеваемость за 3 месяца")
    def student_progress_last_3_months(self, obj: User):
        if obj.role != User.Role.STUDENT:
            return "Только для роли Студент"

        start_date = timezone.now() - timedelta(days=90)
        records = AttendanceRecord.objects.filter(student=obj, lesson__starts_at__gte=start_date).select_related(
            "lesson__group"
        )
        if not records.exists():
            return "За последние 3 месяца данных нет"

        total = records.count()
        present = records.filter(status=AttendanceRecord.Status.PRESENT).count()
        late = records.filter(status=AttendanceRecord.Status.LATE).count()
        absent = records.filter(status=AttendanceRecord.Status.ABSENT).count()

        activity_sum = records.aggregate(total=Sum("activity_score"))["total"] or 0
        homework_done = records.filter(homework_score__gt=0).count()
        exam_avg = records.aggregate(total=Sum("exam_score"))["total"] or 0
        exam_avg = round(exam_avg / total, 2) if total else 0

        attendance_percent = round(((present + late) / total) * 100, 2) if total else 0
        return format_html(
            "Посещаемость: <strong>{}%</strong> (присутствовал: {}, опоздал: {}, отсутствовал: {})<br/>"
            "Активность: <strong>{}/{}</strong><br/>"
            "ДЗ: <strong>{}/{}</strong><br/>"
            "Экзамен (средний): <strong>{}/100</strong>",
            attendance_percent,
            present,
            late,
            absent,
            activity_sum,
            total * 10,
            homework_done,
            total,
            exam_avg,
        )

    @admin.display(description="Нагрузка учителя")
    def teacher_workload_summary(self, obj: User):
        if obj.role != User.Role.TEACHER:
            return "Только для роли Учитель"

        groups = StudyGroup.objects.filter(teacher=obj, is_active=True).prefetch_related("schedules", "enrollments")
        if not groups.exists():
            return "Учителю пока не назначены группы"

        group_count = groups.count()
        student_count = (
            GroupEnrollment.objects.filter(group__in=groups, is_active=True).values("student").distinct().count()
        )

        rows = []
        for group in groups:
            schedule_text = ", ".join(
                f"{schedule.get_day_of_week_display()} {schedule.start_time:%H:%M}-{schedule.end_time:%H:%M}"
                for schedule in group.schedules.all()
            ) or "Нет расписания"
            rows.append(f"<li>{group.subject.name} / {group.name}: {schedule_text}</li>")

        return format_html(
            "Групп: <strong>{}</strong>, учеников: <strong>{}</strong><ul>{}</ul>",
            group_count,
            student_count,
            "".join(rows),
        )

    @admin.display(description="Экспорт")
    def export_user_pdf_button(self, obj: User):
        if not obj.pk:
            return "Сохраните пользователя, чтобы экспортировать"
        url = reverse("admin:crm_user_export_pdf", args=[obj.pk])
        return format_html('<a class="button" href="{}">Экспорт в PDF (A4)</a>', url)

    def get_urls(self):
        urls = super().get_urls()
        custom_urls = [
            path(
                "<path:object_id>/export-pdf/",
                self.admin_site.admin_view(self.export_pdf_view),
                name="crm_user_export_pdf",
            )
        ]
        return custom_urls + urls

    def export_pdf_view(self, request: HttpRequest, object_id: str) -> HttpResponse:
        obj = self.get_object(request, object_id)
        if not obj or not self.has_view_permission(request, obj):
            raise PermissionDenied("Нет доступа к экспорту этого пользователя.")
        if not canvas or not A4:
            raise ValidationError("Библиотека reportlab не установлена.")

        buffer = BytesIO()
        pdf = canvas.Canvas(buffer, pagesize=A4)
        width, height = A4

        y = height - 40
        pdf.setFont("Helvetica-Bold", 14)
        pdf.drawString(40, y, "Карточка пользователя")
        y -= 24

        pdf.setFont("Helvetica", 10)
        lines = [
            f"Логин: {obj.username}",
            f"ФИО: {obj.full_name}",
            f"Роль: {obj.get_role_display()}",
            f"Телефон: {obj.phone_student or '-'}",
            f"Телефон родителя: {obj.phone_parent or '-'}",
            f"Адрес: {obj.address or '-'}",
            f"Школа/смена: {obj.school_name or '-'} / {obj.get_school_shift_display() if obj.school_shift else '-'}",
            f"Уровень теста: {obj.get_placement_level_display() if obj.placement_level else 'Не определен'}",
            f"Балл теста: {obj.placement_score_percent if obj.placement_score_percent is not None else '-'}",
        ]

        for line in lines:
            pdf.drawString(40, y, line)
            y -= 16
            if y < 80:
                pdf.showPage()
                pdf.setFont("Helvetica", 10)
                y = height - 40

        if obj.role == User.Role.STUDENT:
            pdf.setFont("Helvetica-Bold", 11)
            pdf.drawString(40, y, "Группы студента:")
            y -= 18
            pdf.setFont("Helvetica", 10)
            enrollments = GroupEnrollment.objects.filter(student=obj, is_active=True).select_related("group__subject")
            for enrollment in enrollments:
                pdf.drawString(50, y, f"- {enrollment.group.subject.name} / {enrollment.group.name}")
                y -= 14
                if y < 80:
                    pdf.showPage()
                    pdf.setFont("Helvetica", 10)
                    y = height - 40

        pdf.showPage()
        pdf.save()
        buffer.seek(0)

        response = HttpResponse(buffer.getvalue(), content_type="application/pdf")
        response["Content-Disposition"] = f'attachment; filename="user_{obj.pk}.pdf"'
        return response


@admin.register(Subject)
class SubjectAdmin(ModelAdmin):
    list_display = ("name", "is_active")
    list_filter = ("is_active",)
    search_fields = ("name",)

    def has_module_permission(self, request):
        return is_super_admin(request.user) or is_center_admin(request.user)

    def has_view_permission(self, request, obj=None):
        return is_super_admin(request.user) or is_center_admin(request.user)

    def has_add_permission(self, request):
        return is_super_admin(request.user) or is_center_admin(request.user)

    def has_change_permission(self, request, obj=None):
        return is_super_admin(request.user) or is_center_admin(request.user)

    def has_delete_permission(self, request, obj=None):
        return is_super_admin(request.user)


class GroupScheduleInline(admin.TabularInline):
    model = GroupSchedule
    extra = 1


@admin.register(StudyGroup)
class StudyGroupAdmin(ModelAdmin):
    list_display = ("name", "subject", "level", "teacher", "cycle_price", "is_active")
    list_filter = ("subject", "is_active")
    search_fields = ("name", "level", "subject__name")
    inlines = [GroupScheduleInline]

    def get_queryset(self, request):
        queryset = super().get_queryset(request)
        if is_teacher(request.user):
            return queryset.filter(teacher=request.user, is_active=True)
        if is_student(request.user):
            return queryset.filter(enrollments__student=request.user, enrollments__is_active=True).distinct()
        return queryset

    def has_module_permission(self, request):
        return (
            is_super_admin(request.user)
            or is_center_admin(request.user)
            or is_teacher(request.user)
            or is_student(request.user)
        )

    def has_view_permission(self, request, obj=None):
        return self.has_module_permission(request)

    def has_add_permission(self, request):
        return is_super_admin(request.user) or is_center_admin(request.user)

    def has_change_permission(self, request, obj=None):
        if is_super_admin(request.user) or is_center_admin(request.user):
            return True
        return False

    def has_delete_permission(self, request, obj=None):
        return is_super_admin(request.user)


@admin.register(GroupEnrollment)
class GroupEnrollmentAdmin(ModelAdmin):
    list_display = ("group", "student", "started_at", "is_active")
    list_filter = ("group", "is_active")
    search_fields = ("group__name", "student__first_name", "student__last_name", "student__username")

    def get_queryset(self, request):
        queryset = super().get_queryset(request)
        if is_teacher(request.user):
            return queryset.filter(group__teacher=request.user)
        if is_student(request.user):
            return queryset.filter(student=request.user)
        return queryset

    def has_module_permission(self, request):
        return (
            is_super_admin(request.user)
            or is_center_admin(request.user)
            or is_teacher(request.user)
            or is_student(request.user)
        )

    def has_view_permission(self, request, obj=None):
        return self.has_module_permission(request)

    def has_add_permission(self, request):
        return is_super_admin(request.user) or is_center_admin(request.user)

    def has_change_permission(self, request, obj=None):
        return is_super_admin(request.user) or is_center_admin(request.user)

    def has_delete_permission(self, request, obj=None):
        return is_super_admin(request.user)


@admin.register(Lesson)
class LessonAdmin(ModelAdmin):
    list_display = ("topic", "group", "starts_at")
    list_filter = ("group", "group__subject")
    search_fields = ("topic", "group__name", "group__subject__name")
    actions = ["export_homework_pdf", "export_homework_docx"]

    def get_queryset(self, request):
        queryset = super().get_queryset(request)
        if is_teacher(request.user):
            return queryset.filter(group__teacher=request.user)
        if is_student(request.user):
            return queryset.filter(group__enrollments__student=request.user, group__enrollments__is_active=True).distinct()
        return queryset

    def has_module_permission(self, request):
        return (
            is_super_admin(request.user)
            or is_center_admin(request.user)
            or is_teacher(request.user)
            or is_student(request.user)
        )

    def has_view_permission(self, request, obj=None):
        return self.has_module_permission(request)

    def has_add_permission(self, request):
        return is_super_admin(request.user) or is_center_admin(request.user) or is_teacher(request.user)

    def has_change_permission(self, request, obj=None):
        if is_super_admin(request.user) or is_center_admin(request.user):
            return True
        if is_teacher(request.user):
            return obj is None or obj.group.teacher_id == request.user.id
        return False

    def has_delete_permission(self, request, obj=None):
        if is_super_admin(request.user) or is_center_admin(request.user):
            return True
        if is_teacher(request.user):
            return obj is not None and obj.group.teacher_id == request.user.id
        return False

    @admin.action(description="Экспорт выбранных заданий в PDF")
    def export_homework_pdf(self, request, queryset):
        if not (is_super_admin(request.user) or is_center_admin(request.user) or is_teacher(request.user)):
            raise PermissionDenied("Нет доступа к экспорту.")
        if not canvas or not A4:
            self.message_user(request, "Для PDF экспорта нужно установить reportlab.", level=messages.ERROR)
            return None

        queryset = queryset.select_related("group__subject", "group").order_by("group__subject__name", "group__name", "starts_at")
        buffer = BytesIO()
        pdf = canvas.Canvas(buffer, pagesize=A4)
        width, height = A4
        y = height - 40

        pdf.setFont("Helvetica-Bold", 13)
        pdf.drawString(40, y, "Задания учителя")
        y -= 24
        pdf.setFont("Helvetica", 10)

        for lesson in queryset:
            lines = [
                f"{lesson.group.subject.name} / {lesson.group.name}",
                f"Урок: {lesson.topic}",
                f"Дата: {lesson.starts_at:%Y-%m-%d %H:%M}",
                f"Домашнее задание: {lesson.homework or '-'}",
            ]
            for line in lines:
                pdf.drawString(40, y, line[:115])
                y -= 14
                if y < 80:
                    pdf.showPage()
                    pdf.setFont("Helvetica", 10)
                    y = height - 40
            y -= 8

        pdf.showPage()
        pdf.save()
        buffer.seek(0)

        response = HttpResponse(buffer.getvalue(), content_type="application/pdf")
        response["Content-Disposition"] = 'attachment; filename="teacher_homework.pdf"'
        return response

    @admin.action(description="Экспорт выбранных заданий в Word")
    def export_homework_docx(self, request, queryset):
        if not (is_super_admin(request.user) or is_center_admin(request.user) or is_teacher(request.user)):
            raise PermissionDenied("Нет доступа к экспорту.")
        if not Document:
            self.message_user(request, "Для Word экспорта нужно установить python-docx.", level=messages.ERROR)
            return None

        queryset = queryset.select_related("group__subject", "group").order_by("group__subject__name", "group__name", "starts_at")
        document = Document()
        document.add_heading("Задания учителя", level=1)

        for lesson in queryset:
            document.add_heading(f"{lesson.group.subject.name} / {lesson.group.name}", level=2)
            document.add_paragraph(f"Урок: {lesson.topic}")
            document.add_paragraph(f"Дата: {lesson.starts_at:%Y-%m-%d %H:%M}")
            document.add_paragraph(f"Домашнее задание: {lesson.homework or '-'}")

        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)

        response = HttpResponse(
            buffer.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response["Content-Disposition"] = 'attachment; filename="teacher_homework.docx"'
        return response


@admin.register(AttendanceRecord)
class AttendanceRecordAdmin(ModelAdmin):
    list_display = ("lesson", "student", "status", "activity_score", "homework_score", "exam_score")
    list_filter = ("status", "lesson__group", "lesson__group__subject")
    search_fields = ("student__first_name", "student__last_name", "lesson__group__name")
    actions = ["export_attendance_xlsx"]

    def get_queryset(self, request):
        queryset = super().get_queryset(request)
        if is_teacher(request.user):
            return queryset.filter(lesson__group__teacher=request.user)
        if is_student(request.user):
            return queryset.filter(student=request.user)
        return queryset

    def has_module_permission(self, request):
        return (
            is_super_admin(request.user)
            or is_center_admin(request.user)
            or is_teacher(request.user)
            or is_student(request.user)
        )

    def has_view_permission(self, request, obj=None):
        return self.has_module_permission(request)

    def has_add_permission(self, request):
        return is_super_admin(request.user) or is_center_admin(request.user) or is_teacher(request.user)

    def has_change_permission(self, request, obj=None):
        if is_super_admin(request.user) or is_center_admin(request.user):
            return True
        if is_teacher(request.user):
            return obj is None or obj.lesson.group.teacher_id == request.user.id
        if is_student(request.user):
            return obj is None or obj.student_id == request.user.id
        return False

    def has_delete_permission(self, request, obj=None):
        if is_super_admin(request.user) or is_center_admin(request.user):
            return True
        if is_teacher(request.user):
            return obj is not None and obj.lesson.group.teacher_id == request.user.id
        return False

    def get_readonly_fields(self, request, obj=None):
        if is_student(request.user):
            return [field.name for field in self.model._meta.fields]
        return super().get_readonly_fields(request, obj)

    @admin.action(description="Экспорт посещаемости в Excel")
    def export_attendance_xlsx(self, request, queryset):
        if not Workbook:
            self.message_user(request, "Для Excel экспорта нужно установить openpyxl.", level=messages.ERROR)
            return None

        queryset = queryset.select_related("lesson__group__subject", "lesson__group", "student").order_by(
            "lesson__group__subject__name", "lesson__group__name", "lesson__starts_at", "student__last_name"
        )

        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Attendance"
        sheet.append(
            [
                "Предмет",
                "Группа",
                "Дата урока",
                "Ученик",
                "Статус",
                "Активность",
                "ДЗ",
                "Тест",
                "Комментарий",
            ]
        )

        for row in queryset:
            sheet.append(
                [
                    row.lesson.group.subject.name,
                    row.lesson.group.name,
                    row.lesson.starts_at.strftime("%Y-%m-%d %H:%M"),
                    row.student.full_name,
                    row.get_status_display(),
                    row.activity_score,
                    row.homework_score,
                    row.exam_score,
                    row.comment,
                ]
            )

        response = HttpResponse(
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
        response["Content-Disposition"] = 'attachment; filename="attendance_export.xlsx"'
        workbook.save(response)
        return response


@admin.register(BillingCycle)
class BillingCycleAdmin(ModelAdmin):
    list_display = ("group", "number", "opened_at", "lesson_start_number", "lesson_end_number")
    list_filter = ("group__subject", "group")

    def has_module_permission(self, request):
        return is_super_admin(request.user) or is_center_admin(request.user)

    def has_view_permission(self, request, obj=None):
        return is_super_admin(request.user) or is_center_admin(request.user)

    def has_add_permission(self, request):
        return is_super_admin(request.user) or is_center_admin(request.user)

    def has_change_permission(self, request, obj=None):
        return is_super_admin(request.user) or is_center_admin(request.user)

    def has_delete_permission(self, request, obj=None):
        return is_super_admin(request.user)


@admin.register(Payment)
class PaymentAdmin(ModelAdmin):
    list_display = (
        "student",
        "cycle",
        "amount_due",
        "amount_paid",
        "amount_remaining",
        "status_badge",
    )
    list_filter = ("status", "cycle__group__subject", "cycle__group")
    search_fields = ("student__first_name", "student__last_name", "student__username", "cycle__group__name")

    def get_queryset(self, request):
        queryset = super().get_queryset(request)
        if is_teacher(request.user):
            return queryset.none()
        if is_student(request.user):
            return queryset.filter(student=request.user)
        return queryset

    def has_module_permission(self, request):
        return is_super_admin(request.user) or is_center_admin(request.user) or is_student(request.user)

    def has_view_permission(self, request, obj=None):
        if is_super_admin(request.user) or is_center_admin(request.user):
            return True
        if is_student(request.user):
            return obj is None or obj.student_id == request.user.id
        return False

    def has_add_permission(self, request):
        return is_super_admin(request.user) or is_center_admin(request.user)

    def has_change_permission(self, request, obj=None):
        if is_super_admin(request.user) or is_center_admin(request.user):
            return True
        if is_student(request.user):
            return obj is None or obj.student_id == request.user.id
        return False

    def has_delete_permission(self, request, obj=None):
        return is_super_admin(request.user) or is_center_admin(request.user)

    def get_readonly_fields(self, request, obj=None):
        if is_student(request.user):
            return [field.name for field in self.model._meta.fields] + ["amount_remaining"]
        return super().get_readonly_fields(request, obj)

    @admin.display(description="Статус")
    def status_badge(self, obj: Payment):
        colors = {
            Payment.Status.PAID: "#18794e",
            Payment.Status.PARTIAL: "#b86900",
            Payment.Status.UNPAID: "#b42318",
        }
        color = colors.get(obj.status, "#344054")
        return format_html(
            '<span style="padding:3px 8px;border-radius:999px;background:{};color:white;">{}</span>',
            color,
            obj.get_status_display(),
        )


@admin.register(UserExitRecord)
class UserExitRecordAdmin(ModelAdmin):
    form = UserExitRecordForm
    list_display = ("user", "removed_by", "exited_at")
    search_fields = ("user__first_name", "user__last_name", "user__username", "reason")

    def get_queryset(self, request):
        queryset = super().get_queryset(request)
        if is_super_admin(request.user):
            return queryset
        if is_center_admin(request.user):
            return queryset.none()
        return queryset.none()

    def has_module_permission(self, request):
        return is_super_admin(request.user) or is_center_admin(request.user)

    def has_view_permission(self, request, obj=None):
        return is_super_admin(request.user)

    def has_add_permission(self, request):
        return is_super_admin(request.user) or is_center_admin(request.user)

    def has_change_permission(self, request, obj=None):
        return is_super_admin(request.user)

    def has_delete_permission(self, request, obj=None):
        return is_super_admin(request.user)

    def save_model(self, request, obj, form, change):
        if not change:
            obj.removed_by = request.user
        super().save_model(request, obj, form, change)


@admin.register(PlacementTest)
class PlacementTestAdmin(ModelAdmin):
    list_display = ("title", "target_role", "question_count", "is_active")
    list_filter = ("target_role", "is_active")

    def has_module_permission(self, request):
        return is_super_admin(request.user) or is_center_admin(request.user)

    def has_view_permission(self, request, obj=None):
        return is_super_admin(request.user) or is_center_admin(request.user)

    def has_add_permission(self, request):
        return is_super_admin(request.user) or is_center_admin(request.user)

    def has_change_permission(self, request, obj=None):
        return is_super_admin(request.user) or is_center_admin(request.user)

    def has_delete_permission(self, request, obj=None):
        return is_super_admin(request.user)


@admin.register(TestAttempt)
class TestAttemptAdmin(ModelAdmin):
    list_display = ("user", "test", "correct_answers", "total_questions", "score_percent", "level", "taken_at")
    list_filter = ("test__target_role", "level", "is_final")
    search_fields = ("user__first_name", "user__last_name", "user__username", "test__title")

    def get_queryset(self, request):
        queryset = super().get_queryset(request)
        if is_teacher(request.user) or is_student(request.user):
            return queryset.filter(user=request.user)
        return queryset

    def has_module_permission(self, request):
        return (
            is_super_admin(request.user)
            or is_center_admin(request.user)
            or is_teacher(request.user)
            or is_student(request.user)
        )

    def has_view_permission(self, request, obj=None):
        if is_super_admin(request.user) or is_center_admin(request.user):
            return True
        if is_teacher(request.user) or is_student(request.user):
            return obj is None or obj.user_id == request.user.id
        return False

    def has_add_permission(self, request):
        return is_super_admin(request.user) or is_center_admin(request.user)

    def has_change_permission(self, request, obj=None):
        return is_super_admin(request.user) or is_center_admin(request.user)

    def has_delete_permission(self, request, obj=None):
        return is_super_admin(request.user)


original_get_urls = admin.site.get_urls


def custom_admin_urls():
    urls = original_get_urls()
    custom_urls = [
        path(
            "crm/reports/teacher-performance/",
            admin.site.admin_view(teacher_performance_view),
            name="crm_teacher_performance",
        ),
        path(
            "crm/reports/cycle-finance-summary/",
            admin.site.admin_view(cycle_finance_summary_view),
            name="crm_cycle_finance_summary",
        ),
    ]
    return custom_urls + urls


admin.site.get_urls = custom_admin_urls
admin.site.site_header = "Edora CRM"
admin.site.site_title = "Edora CRM"
admin.site.index_title = "Управление учебным центром"
